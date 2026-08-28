"""维权行动方案导出为 Word + 引用已有官方空白模板（不再代填新文书）。

设计原则：不调用 LLM 代填任何新文书，也不编造“民事起诉状”式草稿。用户请求
“生成文书”时，导出的是 node_conclude 已经产出的最终维权行动方案（messages 中
最后一条 AIMessage 内容），并指向一份可下载填写的已有官方空白模板作为参考。
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.agents.legal_guide.document_templates import (
    OfficialDocumentTemplate,
    select_related_official_template,
)

PLAN_WORD_DOC_TYPE = "维权行动方案（Word 版）"
PLAN_WORD_FILENAME = "维权行动方案_法护通.docx"


@dataclass(slots=True)
class GeneratedLegalDocument:
    doc_type: str
    text: str
    docx_bytes: bytes
    filename: str
    missing_fields: list[str]
    official_template: OfficialDocumentTemplate | None = None
    related_official_template: OfficialDocumentTemplate | None = None


def export_plan_word(
    *,
    legal_domain: str,
    plan_text: str,
    confirmed_issues: list[str] | None = None,
    collected_facts: list[str] | None = None,
) -> GeneratedLegalDocument:
    """把已生成的维权行动方案导出为可编辑 Word，并引用已有官方空白模板。

    - plan_text 即 node_conclude 产出的最终方案（最后一条 AIMessage 内容）；
    - 复用 render_legal_docx 渲染为 DOCX，无 LLM 调用、不代填任何新文书；
    - 通过 select_related_official_template 挑一个已有官方模板作为参考引用；
      若当前领域没有匹配的官方模板（如刑事/治安领域暂无官方示范文本），
      则不编造模板引用，仅交付方案 Word 本身。
    """
    plan_text = str(plan_text or "").strip() or "（当前尚未生成可导出的维权行动方案）"
    related_official_template = select_related_official_template(
        legal_domain,
        list(confirmed_issues or []),
        list(collected_facts or []),
    )
    return GeneratedLegalDocument(
        doc_type=PLAN_WORD_DOC_TYPE,
        text=plan_text,
        docx_bytes=render_legal_docx(
            PLAN_WORD_DOC_TYPE,
            plan_text,
            related_official_template,
        ),
        filename=PLAN_WORD_FILENAME,
        missing_fields=[],
        official_template=None,
        related_official_template=related_official_template,
    )


def _clean_line(line: str) -> str:
    value = line.strip()
    value = re.sub(r"^#{1,6}\s*", "", value)
    value = value.replace("**", "")
    return value


def _set_run_font(run, name: str, size: int, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def render_legal_docx(
    doc_type: str,
    doc_text: str,
    official_template: OfficialDocumentTemplate | None = None,
) -> bytes:
    """Render an editable DOCX while keeping source and system status explicit."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    lines = [_clean_line(line) for line in doc_text.splitlines()]
    lines = [line for line in lines if line]
    title_written = False
    for line in lines:
        if not title_written:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(paragraph.add_run(line), "黑体", 18, True)
            title_written = True
            continue
        is_heading = line.rstrip("：") in {
            "当事人信息",
            "诉讼请求",
            "申请事项",
            "投诉请求",
            "事实与理由",
            "事实经过",
            "证据清单",
            "随附证据",
            "对纠纷解决方式的意愿",
        }
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0 if is_heading else 0.74)
        paragraph.paragraph_format.line_spacing = 1.5
        _set_run_font(
            paragraph.add_run(line),
            "黑体" if is_heading else "宋体",
            12,
            is_heading,
        )

    document.add_paragraph()
    source_heading = document.add_paragraph()
    _set_run_font(source_heading.add_run("模板来源与文件性质"), "黑体", 11, True)
    if official_template:
        collection = official_template.collection
        source_text = (
            f"可参考的官方空白模板：{collection.title}（{collection.document_no}），"
            f"发布机关：{'、'.join(collection.issuers)}，"
            f"自 {collection.effective_at} 起推广使用。官方原文："
            f"{collection.source_page_url}"
        )
    else:
        source_text = (
            "当前领域尚未匹配到可参考的官方空白模板，本文件为用户已生成的维权行动方案。"
        )
    source_paragraph = document.add_paragraph()
    source_run = source_paragraph.add_run(source_text)
    _set_run_font(source_run, "宋体", 9)
    source_run.font.color.rgb = RGBColor(89, 89, 89)

    disclaimer = document.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "本文件为系统根据用户提供的信息生成的可编辑参考稿，非人民法院、司法行政机关或其他发布机关出具。"
    )
    _set_run_font(disclaimer_run, "宋体", 9, True)
    disclaimer_run.font.color.rgb = RGBColor(192, 57, 43)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
