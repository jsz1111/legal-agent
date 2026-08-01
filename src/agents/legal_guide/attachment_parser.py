"""Bounded, auditable extraction for user-supplied document attachments."""
from __future__ import annotations

import hashlib
import io
from pathlib import Path


SUPPORTED_DOCUMENT_SUFFIXES = {".txt", ".docx", ".pdf"}
MAX_ATTACHMENT_BYTES = 10 * 1024 * 1024
MAX_ATTACHMENT_TEXT = 8_000


def extract_document_bytes(
    filename: str,
    content: bytes,
    *,
    max_bytes: int = MAX_ATTACHMENT_BYTES,
    max_text: int = MAX_ATTACHMENT_TEXT,
) -> dict:
    """Extract bounded text without retaining the uploaded document."""

    safe_name = Path(filename or "未命名文件").name[:180]
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
        raise ValueError("仅支持 PDF、DOCX 和 TXT 文档")
    if not content:
        raise ValueError("文件内容为空")
    if len(content) > max_bytes:
        raise ValueError("文件超过10MB，请压缩后再上传")

    source_form = "native_electronic"
    text = ""
    if suffix == ".txt":
        text = content.decode("utf-8", errors="replace")
    elif suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(content))
        blocks = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
        text = "\n".join(blocks)
    else:
        from pypdf import PdfReader

        source_form = "exported_file"
        reader = PdfReader(io.BytesIO(content))
        blocks: list[str] = []
        for page in reader.pages[:30]:
            blocks.append(page.extract_text() or "")
            if sum(len(item) for item in blocks) >= max_text:
                break
        text = "\n".join(blocks)

    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    scan_warning = False
    if not normalized:
        scan_warning = True
        normalized = (
            "未提取到可读文字；该文件可能是扫描件。"
            "请改传清晰图片，或在案情中手动说明关键内容。"
        )
    truncated = len(normalized) > max_text
    extracted_text = normalized[:max_text]
    digest = hashlib.sha256(content).hexdigest()
    truncation_note = (
        "\n（内容较长，本轮仅提取前部文字。）"
        if truncated
        else ""
    )
    evidence_block = (
        "【文档证据补充（程序提取，需与原文件核对）】\n"
        f"文件：{safe_name}\n"
        f"来源形式：{source_form}\n"
        f"原文件 SHA-256：{digest}\n"
        "【提取文字】\n"
        f"{extracted_text}{truncation_note}"
    )
    return {
        "filename": safe_name,
        "text": extracted_text,
        "preview": extracted_text[:500],
        "sha256": digest,
        "source_form": source_form,
        "truncated": truncated,
        "scan_warning": scan_warning,
        "evidence_block": evidence_block,
        "retained": False,
        "size_bytes": len(content),
    }
