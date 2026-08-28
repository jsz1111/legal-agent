from __future__ import annotations

import hashlib
import io

from docx import Document
from pypdf import PdfReader

from src.agents.legal_guide.doc_generator import (
    export_plan_word,
    render_legal_docx,
)
from src.agents.legal_guide.formatters import requested_doc_type
from src.agents.legal_guide.document_templates import (
    list_official_templates,
    select_official_template,
    select_related_official_template,
)


def test_official_manifest_files_and_hashes_are_valid():
    templates = list_official_templates()
    assert len(templates) == 8

    for template in templates:
        assert template.blank_pdf_path.is_file()
        assert template.evidence_items
        digest = hashlib.sha256(template.blank_pdf_path.read_bytes()).hexdigest().upper()
        assert digest == template.blank_pdf_sha256
        reader = PdfReader(template.blank_pdf_path)
        assert len(reader.pages) >= 4
        first_page = "".join((reader.pages[0].extract_text() or "").split())
        assert "民事起诉状" in first_page


def test_template_selector_uses_case_type_and_procedural_stage():
    lease = select_official_template(
        "contracts_property_housing",
        "房东拒不退还租房押金",
    )
    assert lease and lease.template_id == "spc_2025_house_lease_complaint"

    consumer = select_official_template(
        "consumer_market",
        "网购商品存在质量问题，商家拒绝退款",
    )
    assert consumer and consumer.template_id == "spc_2025_sales_contract_complaint"

    assert select_official_template(
        "labor_social_security",
        "公司拖欠工资，尚未申请劳动仲裁",
    ) is None
    related_labor = select_related_official_template(
        "labor_social_security",
        "公司拖欠工资，尚未申请劳动仲裁",
    )
    assert related_labor
    assert related_labor.template_id == "spc_2025_labor_dispute_complaint"

    labor_litigation = select_official_template(
        "labor_social_security",
        "不服劳动仲裁裁决，准备向法院起诉",
    )
    assert labor_litigation
    assert labor_litigation.template_id == "spc_2025_labor_dispute_complaint"


def test_rendered_docx_is_editable_and_labels_source():
    template = select_official_template(
        "traffic_personal_injury",
        "交通事故责任认定后准备起诉",
    )
    assert template
    payload = render_legal_docx(
        template.title,
        "民事起诉状（机动车交通事故责任纠纷）\n"
        "当事人信息：\n原告：【请填写姓名】\n"
        "诉讼请求：\n1. 赔偿医疗费【请填写金额】元。",
        template,
    )
    assert payload.startswith(b"PK")

    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "法〔2025〕82号" in text
    assert "系统根据用户提供的信息生成" in text
    assert "非人民法院" in text


def test_export_plan_word_exports_plan_with_reference_template_for_civil_domain():
    result = export_plan_word(
        legal_domain="labor_social_security",
        plan_text="**维权行动方案**\n1. 先向劳动监察投诉。\n2. 再申请劳动仲裁。",
        confirmed_issues=["拖欠劳动报酬"],
        collected_facts=["公司拖欠工资三个月"],
    )

    assert result.doc_type == "维权行动方案（Word 版）"
    assert result.filename == "维权行动方案_法护通.docx"
    assert result.missing_fields == []
    # 方案 Word 不代填新文书：不绑定“精确”模板，只附带同领域官方模板作为参考。
    assert result.official_template is None
    assert result.related_official_template
    assert result.related_official_template.template_id == "spc_2025_labor_dispute_complaint"
    assert result.docx_bytes.startswith(b"PK")

    document = Document(io.BytesIO(result.docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "维权行动方案" in text
    assert "劳动仲裁" in text
    assert "官方空白模板" in text


def test_export_plan_word_for_criminal_domain_does_not_invent_template():
    # 刑事/治安领域没有官方空白模板，导出方案 Word 时不编造模板引用。
    result = export_plan_word(
        legal_domain="criminal_public_security",
        plan_text="**维权行动方案**\n1. 先报警受案。\n2. 申请伤情鉴定。",
        confirmed_issues=["故意伤害"],
        collected_facts=["被人打伤"],
    )

    assert result.doc_type == "维权行动方案（Word 版）"
    assert result.related_official_template is None
    assert result.missing_fields == []
    assert result.docx_bytes.startswith(b"PK")

    document = Document(io.BytesIO(result.docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "当前领域尚未匹配到可参考的官方空白模板" in text


def test_generic_arbitration_request_uses_domain_default_to_disambiguate():
    assert requested_doc_type("生成仲裁申请书", "劳动仲裁申请书") == "劳动仲裁申请书"
    assert requested_doc_type("生成仲裁申请书", "仲裁申请书") == "仲裁申请书"
